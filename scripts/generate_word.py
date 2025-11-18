import os
from docx import Document
from docx.shared import Pt

# Diretórios e arquivos a incluir
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
INCLUDE_PATHS = [
    os.path.join(ROOT, 'README.md'),
    os.path.join(ROOT, 'pom.xml'),
    os.path.join(ROOT, 'src', 'main', 'java'),
    os.path.join(ROOT, 'src', 'main', 'resources'),
]

EXTS_CODE = {'.java', '.xml', '.md', '.properties'}


def is_code_file(path):
    _, ext = os.path.splitext(path)
    return ext.lower() in EXTS_CODE


def add_file_to_doc(doc, filepath, relpath):
    doc.add_heading(relpath, level=2)
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
    except Exception:
        try:
            with open(filepath, 'r', encoding='latin-1') as f:
                content = f.read()
        except Exception as e:
            doc.add_paragraph(f'Não foi possível ler o arquivo: {e}')
            return

    # Adiciona conteúdo como bloco de código
    p = doc.add_paragraph()
    run = p.add_run(content)
    font = run.font
    font.name = 'Courier New'
    font.size = Pt(8)


def collect_files():
    files = []
    for path in INCLUDE_PATHS:
        if os.path.isfile(path):
            files.append((path, os.path.relpath(path, ROOT)))
        elif os.path.isdir(path):
            for root, _, filenames in os.walk(path):
                for name in sorted(filenames):
                    fp = os.path.join(root, name)
                    if is_code_file(fp):
                        files.append((fp, os.path.relpath(fp, ROOT)))
    return files


def main():
    doc = Document()
    doc.add_heading('Projeto - Código Fonte', level=1)

    files = collect_files()
    if not files:
        doc.add_paragraph('Nenhum arquivo selecionado para inclusão.')
    for fp, rel in files:
        add_file_to_doc(doc, fp, rel)

    out = os.path.join(ROOT, 'Projeto_codigo.docx')
    doc.save(out)
    print('Arquivo gerado:', out)


if __name__ == '__main__':
    main()
